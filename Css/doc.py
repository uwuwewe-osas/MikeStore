from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def crear_documento_sgsi():
    # Crear documento
    doc = Document()
    
    # Configurar márgenes (2.5 cm como en el original)
    seccion = doc.sections[0]
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)
    
    # ========== PORTADA ==========
    # Título principal
    titulo1 = doc.add_paragraph()
    titulo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo1.add_run("MANUAL DEL SISTEMA DE GESTIÓN DE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    
    titulo2 = doc.add_paragraph()
    titulo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo2.add_run("SEGURIDAD DE LA INFORMACIÓN")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    
    titulo3 = doc.add_paragraph()
    titulo3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo3.add_run("(SGSI)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espacio
    
    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("PARA EL SERVIDOR WEB (FÍSICO Y VIRTUAL)")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Tabla de información
    tabla_info = doc.add_table(rows=4, cols=2)
    tabla_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configurar ancho de columnas
    for row in tabla_info.rows:
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(3)
    
    # Llenar tabla
    datos = [
        ("Clasificación:", "Uso Interno Restringido"),
        ("Versión:", "1.0"),
        ("Área responsable:", "Tecnologías de la Información"),
        ("Año:", "2026")
    ]
    
    for i, (etiqueta, valor) in enumerate(datos):
        p_etiqueta = tabla_info.cell(i, 0).paragraphs[0]
        p_etiqueta.add_run(etiqueta).bold = True
        
        p_valor = tabla_info.cell(i, 1).paragraphs[0]
        p_valor.add_run(valor)
    
    doc.add_page_break()
    
    # ========== CONTROL DE VERSIONES ==========
    titulo = doc.add_heading('CONTROL DE VERSIONES Y APROBACIÓN', level=1)
    
    # Tabla de versiones
    tabla_versiones = doc.add_table(rows=2, cols=4)
    tabla_versiones.style = 'Table Grid'
    
    # Encabezados
    encabezados = ['Versión', 'Fecha', 'Descripción', 'Elaborado/Aprobado']
    for i, encabezado in enumerate(encabezados):
        celda = tabla_versiones.cell(0, i)
        p = celda.paragraphs[0]
        p.add_run(encabezado).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos de la versión
    datos_version = ['1.0', '2026', 
                     'Versión inicial del manual SGSI para servidor web físico y virtual. Cubre niveles 1 a 5 completos.',
                     'Área de TI']
    for i, dato in enumerate(datos_version):
        celda = tabla_versiones.cell(1, i)
        celda.paragraphs[0].add_run(dato)
    
    doc.add_page_break()
    
    # ========== 1. INTRODUCCIÓN ==========
    doc.add_heading('1. INTRODUCCIÓN', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run("La organización reconoce que el servidor web constituye un activo estratégico dentro de su infraestructura tecnológica, al permitir la publicación de servicios institucionales, el intercambio de información con usuarios internos y externos, y el procesamiento de datos relevantes para la operación del negocio. Debido a su naturaleza y nivel de exposición, tanto el servidor web físico como el servidor web virtual se encuentran sujetos a riesgos asociados a accesos no autorizados, explotación de vulnerabilidades, interrupciones del servicio, pérdida de información y otros incidentes que podrían afectar la continuidad operativa y la reputación institucional.")
    
    doc.add_paragraph()
    
    intro2 = doc.add_paragraph()
    intro2.add_run("En ese contexto, resulta necesario establecer un marco formal que defina los lineamientos de seguridad aplicables a la instalación, configuración, administración y mantenimiento de dichos servidores. El presente documento forma parte del Sistema de Gestión de Seguridad de la Información (SGSI) de la organización y establece las disposiciones que deberán cumplirse para garantizar la adecuada protección de los activos de información relacionados con el servicio web. Su implementación contribuye a fortalecer la postura de seguridad tecnológica, reducir la superficie de exposición ante amenazas y asegurar la continuidad del servicio bajo criterios de control y supervisión definidos.")
    
    # 1.1 Objetivo
    doc.add_heading('1.1 Objetivo', level=2)
    objetivo = doc.add_paragraph()
    objetivo.add_run("El objetivo del presente documento es establecer las políticas y lineamientos de seguridad aplicables al servidor web físico y al servidor web virtual de la organización, con la finalidad de proteger la confidencialidad, integridad y disponibilidad de la información gestionada a través de dichos entornos. Asimismo, busca definir criterios claros para su administración segura, establecer responsabilidades en su gestión y asegurar que su configuración y operación se realicen conforme a estándares y buenas prácticas reconocidas en materia de seguridad de la información.")
    
    # 1.2 Alcance
    doc.add_heading('1.2 Alcance', level=2)
    alcance = doc.add_paragraph()
    alcance.add_run("El presente documento es de aplicación obligatoria para el servidor web físico ubicado en las instalaciones de la organización y para el servidor web virtual implementado sobre infraestructura de virtualización interna o alojado en un proveedor externo de servicios. El alcance comprende los sistemas operativos instalados, los servicios web desplegados, las bases de datos asociadas, los mecanismos de autenticación y control de acceso, así como los dispositivos y configuraciones de red vinculados a su funcionamiento.")
    
    alcance2 = doc.add_paragraph()
    alcance2.add_run("Asimismo, incluye al personal del área de Tecnologías de la Información, administradores de sistemas, proveedores externos y cualquier persona que cuente con acceso físico o lógico a los servidores mencionados. Se encuentran dentro del alcance las actividades relacionadas con su instalación, configuración inicial, actualización, monitoreo, respaldo, mantenimiento y recuperación ante incidentes.")
    
    # 1.3 Justificación
    doc.add_heading('1.3 Justificación', level=2)
    just = doc.add_paragraph()
    just.add_run("La implementación de un conjunto estructurado de lineamientos de seguridad para el servidor web físico y virtual responde a la necesidad de proteger un activo tecnológico crítico para la organización. El servidor web constituye un punto de exposición permanente hacia redes internas y externas, lo que incrementa la probabilidad de incidentes asociados a accesos no autorizados, explotación de vulnerabilidades, alteración de información o interrupciones del servicio.")
    
    just2 = doc.add_paragraph()
    just2.add_run("La ausencia de políticas formales y controles definidos puede generar configuraciones inconsistentes, prácticas administrativas no estandarizadas y una mayor vulnerabilidad frente a amenazas internas y externas. En ese contexto, el presente manual establece un marco de referencia que permite definir responsabilidades, aplicar controles adecuados y reducir los riesgos asociados a la operación del servidor web. Su implementación favorece la continuidad operativa, la protección de la reputación institucional y la confianza de clientes y usuarios.")
    
    # 1.4 Metodología
    doc.add_heading('1.4 Metodología', level=2)
    metodologia = doc.add_paragraph()
    metodologia.add_run("La elaboración del presente manual se ha desarrollado bajo un enfoque basado en la gestión de riesgos y en la adopción de buenas prácticas internacionales en seguridad de la información. Como primera etapa, se realizó una identificación preliminar de los activos asociados al servidor web físico y virtual, considerando tanto los componentes tecnológicos como los procesos y personas involucradas en su gestión.")
    
    metodologia2 = doc.add_paragraph()
    metodologia2.add_run("Posteriormente, se analizaron las amenazas y vulnerabilidades que podrían afectar dichos activos, evaluando su posible impacto en la operación del servicio web. Sobre la base de este análisis, se definieron políticas y lineamientos orientados a mitigar los riesgos identificados, priorizando aquellos que representan mayor criticidad para la organización. La metodología contempla además la definición de controles preventivos, detectivos y correctivos, así como la revisión periódica del manual con el fin de asegurar su actualización frente a cambios tecnológicos, nuevos escenarios de riesgo o modificaciones en la infraestructura del servidor web.")
    
    doc.add_page_break()
    
    # ========== 2. MARCO NORMATIVO Y REFERENCIAL ==========
    doc.add_heading('2. MARCO NORMATIVO Y REFERENCIAL', level=1)
    
    marco = doc.add_paragraph()
    marco.add_run("El presente manual se fundamenta en estándares internacionales y buenas prácticas reconocidas en materia de seguridad de la información, seguridad de redes y protección de infraestructuras tecnológicas. Estos marcos normativos proporcionan lineamientos estructurados para la identificación, evaluación y tratamiento de riesgos que puedan afectar los activos asociados al servidor web físico y virtual.")
    
    # 2.1 ISO/IEC 27001
    doc.add_heading('2.1 ISO/IEC 27001', level=2)
    iso = doc.add_paragraph()
    iso.add_run("La norma ISO/IEC 27001 establece los requisitos para la implementación, mantenimiento y mejora continua de un Sistema de Gestión de Seguridad de la Información (SGSI). Este estándar define un modelo de gestión basado en la identificación y tratamiento sistemático de riesgos que puedan comprometer la confidencialidad, integridad y disponibilidad de la información. El presente manual adopta los principios de dicho estándar en lo relacionado con:")
    
    # Lista de puntos ISO
    lista_iso = [
        "Enfoque basado en riesgos para la toma de decisiones de seguridad.",
        "Definición y aprobación de políticas de seguridad de la información.",
        "Asignación de roles y responsabilidades formales.",
        "Gestión y protección de activos de información.",
        "Control de accesos físicos y lógicos.",
        "Seguridad física, ambiental y en redes y comunicaciones.",
        "Gestión de incidentes de seguridad y mejora continua del sistema."
    ]
    
    for punto in lista_iso:
        p = doc.add_paragraph(punto, style='List Bullet')
    
    iso2 = doc.add_paragraph()
    iso2.add_run("De manera complementaria, se toma como referencia la norma ISO/IEC 27002, la cual proporciona directrices detalladas para la implementación práctica de controles de seguridad. En el contexto nacional, se considera la adopción de este estándar a través de la correspondiente Norma Técnica Peruana.")
    
    # 2.2 Principios de Seguridad
    doc.add_heading('2.2 Principios de Seguridad de la Información', level=2)
    
    principios = doc.add_paragraph()
    principios.add_run("El presente manual se sustenta en los tres principios fundamentales de la seguridad de la información, que constituyen la base conceptual sobre la cual se estructuran todas las políticas y controles aplicables:")
    
    # Tabla de principios
    tabla_principios = doc.add_table(rows=4, cols=2)
    tabla_principios.style = 'Table Grid'
    
    # Encabezados
    tabla_principios.cell(0, 0).paragraphs[0].add_run("Principio").bold = True
    tabla_principios.cell(0, 1).paragraphs[0].add_run("Descripción").bold = True
    
    # Datos
    datos_principios = [
        ("Confidencialidad", "Garantizar que la información sea accesible únicamente por personas debidamente autorizadas mediante controles de acceso, autenticación y cifrado."),
        ("Integridad", "Proteger la exactitud y completitud de la información frente a modificaciones no autorizadas o accidentales mediante controles de integridad y auditoría."),
        ("Disponibilidad", "Asegurar que los sistemas y la información estén accesibles cuando sean requeridos para la operación del negocio mediante redundancia, respaldos y planes de continuidad.")
    ]
    
    for i, (principio, desc) in enumerate(datos_principios, 1):
        tabla_principios.cell(i, 0).paragraphs[0].add_run(principio)
        tabla_principios.cell(i, 1).paragraphs[0].add_run(desc)
    
    # 2.3 Ciclo PHVA
    doc.add_heading('2.3 Ciclo PHVA', level=2)
    
    phva = doc.add_paragraph()
    phva.add_run("La gestión de la seguridad del servidor web se desarrolla bajo el enfoque de mejora continua conocido como ciclo PHVA (Planificar -- Hacer -- Verificar -- Actuar), el cual forma parte integral del modelo de gestión establecido en la norma ISO/IEC 27001:")
    
    # Tabla PHVA
    tabla_phva = doc.add_table(rows=5, cols=3)
    tabla_phva.style = 'Table Grid'
    
    # Encabezados
    headers = ["Fase", "Sigla", "Descripción"]
    for i, header in enumerate(headers):
        tabla_phva.cell(0, i).paragraphs[0].add_run(header).bold = True
    
    # Datos PHVA
    datos_phva = [
        ("Planificar", "P", "Identificación de activos, evaluación de riesgos y definición de políticas y controles de seguridad."),
        ("Hacer", "H", "Implementación de las políticas, configuraciones técnicas y controles de seguridad establecidos."),
        ("Verificar", "V", "Monitoreo del servidor, revisión de registros, auditorías internas y evaluación del cumplimiento."),
        ("Actuar", "A", "Corrección de desviaciones, tratamiento de incidentes y mejora continua de las medidas implementadas.")
    ]
    
    for i, (fase, sigla, desc) in enumerate(datos_phva, 1):
        tabla_phva.cell(i, 0).paragraphs[0].add_run(fase)
        tabla_phva.cell(i, 1).paragraphs[0].add_run(sigla)
        tabla_phva.cell(i, 2).paragraphs[0].add_run(desc)
    
    doc.add_page_break()
    
    # ========== NIVEL 2 - POLÍTICAS ==========
    # Separador de nivel
    p_nivel = doc.add_paragraph()
    p_nivel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_nivel.add_run("NIVEL 2 -- DEFINICIÓN DE POLÍTICAS DE SEGURIDAD")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # ========== 1. POLÍTICAS PARA SERVIDORES WEB ON-PREMISE ==========
    doc.add_heading('1. POLÍTICAS PARA SERVIDORES WEB ON-PREMISE', level=1)
    
    # 1.1 Política de Seguridad On-Premise
    doc.add_heading('1.1 Política de Seguridad On-Premise', level=2)
    
    # Marco normativo
    marco_normativo = doc.add_paragraph()
    marco_normativo.add_run("Marco Normativo: ").bold = True
    marco_normativo.add_run("ISO/IEC 27001:2022 -- Controles 7.1, 7.2, 7.5 y 7.11 | ISO/IEC 27002 -- Seguridad física y ambiental")
    
    # 1.1.1 Objetivo
    doc.add_heading('1.1.1 Objetivo', level=3)
    obj = doc.add_paragraph()
    obj.add_run("Establecer lineamientos para garantizar la protección on-premise del servidor web frente a accesos no autorizados, daños intencionales, fallas ambientales y riesgos operativos que puedan comprometer la confidencialidad, integridad y disponibilidad de la información.")
    
    # 1.1.2 Control de Acceso al Área del Servidor
    doc.add_heading('1.1.2 Control de Acceso al Área del Servidor', level=3)
    
    items_acceso = [
        "El servidor web deberá ubicarse en un área físicamente restringida y con acceso controlado.",
        "El acceso estará limitado exclusivamente al personal autorizado, con registro documentado.",
        "El área deberá permanecer cerrada con mecanismos de control: llave, tarjeta de proximidad o sistema biométrico equivalente.",
        "Toda visita de personal externo al área del servidor deberá estar supervisada por personal de TI.",
        "No se permitirá almacenamiento de materiales ajenos al servicio en el área del servidor.",
        "Se deberá mantener un registro documentado de todos los accesos físicos, con fecha, hora, nombre y motivo."
    ]
    
    for item in items_acceso:
        p = doc.add_paragraph(item, style='List Number')
    
    # 1.1.3 Protección Eléctrica (UPS y Estabilizador)
    doc.add_heading('1.1.3 Protección Eléctrica (UPS y Estabilizador)', level=3)
    
    items_electricos = [
        "El servidor deberá contar con sistema de alimentación ininterrumpida (UPS) con autonomía suficiente para una parada ordenada.",
        "Se deberá emplear estabilizador o regulador de voltaje para proteger el equipo de variaciones eléctricas.",
        "El sistema eléctrico deberá someterse a mantenimiento preventivo periódico, con registro documentado.",
        "Se deberá establecer un procedimiento documentado ante interrupciones prolongadas de energía.",
        "Se recomienda la conexión a sistema de puesta a tierra certificado."
    ]
    
    for item in items_electricos:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ========== 1.2 Política de Instalación Segura ==========
    doc.add_heading('1.2 Política de Instalación Segura', level=2)
    
    marco_normativo2 = doc.add_paragraph()
    marco_normativo2.add_run("Marco Normativo: ").bold = True
    marco_normativo2.add_run("ISO/IEC 27001:2022 -- Controles 8.8 y 8.9 | ISO/IEC 27002 -- Gestión de vulnerabilidades y configuración segura")
    
    # 1.2.1 Objetivo
    doc.add_heading('1.2.1 Objetivo', level=3)
    obj2 = doc.add_paragraph()
    obj2.add_run("Establecer lineamientos para asegurar que la instalación del sistema operativo y componentes del servidor web se realice bajo criterios de seguridad desde su implementación inicial, reduciendo la superficie de ataque desde el primer momento.")
    
    # 1.2.2 Instalación Mínima del Sistema Operativo
    doc.add_heading('1.2.2 Instalación Mínima del Sistema Operativo', level=3)
    
    items_instalacion = [
        "Se deberá realizar instalación mínima del sistema operativo, sin entorno gráfico.",
        "No se instalarán servicios o aplicaciones no necesarios para la función del servidor web.",
        "Se deberán eliminar o deshabilitar servicios por defecto que no sean requeridos.",
        "Se deberá documentar la configuración base del sistema en el repositorio del SGSI.",
        "Se deberán aplicar todas las actualizaciones de seguridad disponibles antes de entrar en producción."
    ]
    
    for item in items_instalacion:
        p = doc.add_paragraph(item, style='List Number')
    
    # 1.2.3 Configuración Inicial Segura
    doc.add_heading('1.2.3 Configuración Inicial Segura', level=3)
    
    items_config = [
        "Se deberán cambiar todas las credenciales por defecto antes de la puesta en producción.",
        "Se deberán aplicar configuraciones de hardening sobre el sistema operativo y el servidor web.",
        "Se deberá configurar política de auditoría y registro de eventos desde el inicio.",
        "Se deberá validar que el sistema se encuentre completamente actualizado.",
        "Toda configuración inicial deberá quedar documentada con fecha, responsable y versiones."
    ]
    
    for item in items_config:
        p = doc.add_paragraph(item, style='List Number')
    
    # ========== 1.3 Política de Gestión de Accesos ==========
    doc.add_heading('1.3 Política de Gestión de Accesos', level=2)
    
    marco_normativo3 = doc.add_paragraph()
    marco_normativo3.add_run("Marco Normativo: ").bold = True
    marco_normativo3.add_run("ISO/IEC 27001:2022 -- Controles 5.15, 5.16 y 5.17 | ISO/IEC 27002 -- Control de acceso e identidad")
    
    # 1.3.1 Objetivo
    doc.add_heading('1.3.1 Objetivo', level=3)
    obj3 = doc.add_paragraph()
    obj3.add_run("Garantizar que el acceso lógico al servidor web físico sea controlado, trazable y restringido bajo el principio de mínimo privilegio, asegurando que cada usuario tenga acceso únicamente a los recursos necesarios para el desempeño de sus funciones.")
    
    # 1.3.2 Administración de Usuarios
    doc.add_heading('1.3.2 Administración de Usuarios', level=3)
    
    items_usuarios = [
        "Cada administrador deberá contar con una cuenta individual. Se prohíbe el uso compartido de credenciales.",
        "Los privilegios deberán asignarse según las funciones documentadas de cada usuario.",
        "Se deberá revocar el acceso de forma inmediata cuando ya no sea necesario (cese, cambio de área, finalización de contrato).",
        "Se deberá mantener un registro documentado de altas, bajas y modificaciones de cuentas.",
        "Está prohibido el uso de cuentas genéricas como 'admin', 'root' o 'soporte' para acceso cotidiano."
    ]
    
    for item in items_usuarios:
        p = doc.add_paragraph(item, style='List Number')
    
    # Guardar documento
    doc.save('SGSI_Manual_Servidor_Web_Parte1.docx')
    print("Documento creado exitosamente: SGSI_Manual_Servidor_Web_Parte1.docx")

if __name__ == "__main__":
    crear_documento_sgsi()